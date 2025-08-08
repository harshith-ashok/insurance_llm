import aiohttp
import PyPDF2
import io
import docx
import re
from typing import Dict, List, Any
import logging

logger = logging.getLogger(__name__)


class DocumentProcessor:
    def __init__(self):
        self.supported_extensions = ['.pdf', '.docx', '.doc']

    async def process_document(self, document_url: str) -> Dict[str, Any]:
        try:
            content = await self._fetch_document(document_url)
            file_extension = self._get_file_extension(document_url)

            if file_extension == '.pdf':
                return await self._process_pdf(content)
            elif file_extension in ['.docx', '.doc']:
                return await self._process_docx(content)
            else:
                return await self._process_text(content)

        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            raise Exception(f"Failed to process document: {str(e)}")

    async def _fetch_document(self, url: str) -> bytes:
        async with aiohttp.ClientSession() as session:
            async with session.get(url) as response:
                if response.status != 200:
                    raise Exception(
                        f"Failed to fetch document: {response.status}")
                return await response.read()

    def _get_file_extension(self, url: str) -> str:
        url_lower = url.lower()
        for ext in self.supported_extensions:
            if ext in url_lower:
                return ext
        return '.txt'

    async def _process_pdf(self, content: bytes) -> Dict[str, Any]:
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text_content = ""
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"

            clauses = self._extract_clauses(text_content)

            return {
                "type": "pdf",
                "content": text_content,
                "clauses": clauses,
                "metadata": {
                    "pages": len(pdf_reader.pages),
                    "extracted_text_length": len(text_content)
                }
            }
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            raise Exception(f"PDF processing failed: {str(e)}")

    async def _process_docx(self, content: bytes) -> Dict[str, Any]:
        try:
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)

            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"

            clauses = self._extract_clauses(text_content)

            return {
                "type": "docx",
                "content": text_content,
                "clauses": clauses,
                "metadata": {
                    "paragraphs": len(doc.paragraphs),
                    "extracted_text_length": len(text_content)
                }
            }
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            raise Exception(f"DOCX processing failed: {str(e)}")

    async def _process_text(self, content: bytes) -> Dict[str, Any]:
        try:
            text_content = content.decode('utf-8', errors='ignore')
            clauses = self._extract_clauses(text_content)

            return {
                "type": "text",
                "content": text_content,
                "clauses": clauses,
                "metadata": {
                    "extracted_text_length": len(text_content)
                }
            }
        except Exception as e:
            logger.error(f"Error processing text: {str(e)}")
            raise Exception(f"Text processing failed: {str(e)}")

    def _extract_clauses(self, text: str) -> List[Dict[str, Any]]:
        clauses = []

        sections = re.split(r'\n\s*\n', text)

        for i, section in enumerate(sections):
            if len(section.strip()) > 50:
                clauses.append({
                    "id": f"clause_{i}",
                    "content": section.strip(),
                    "length": len(section.strip()),
                    "type": self._classify_clause_type(section)
                })

        return clauses

    def _classify_clause_type(self, text: str) -> str:
        text_lower = text.lower()

        if any(word in text_lower for word in ['coverage', 'cover', 'policy']):
            return 'coverage'
        elif any(word in text_lower for word in ['exclusion', 'exclude', 'not covered']):
            return 'exclusion'
        elif any(word in text_lower for word in ['condition', 'requirement', 'must']):
            return 'condition'
        elif any(word in text_lower for word in ['limit', 'maximum', 'cap']):
            return 'limit'
        else:
            return 'general'
